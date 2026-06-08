from __future__ import annotations

import io

from docx import Document
from docx.shared import Pt

from ...models import TranscriptionJob


def _format_duration(seconds: float | None) -> str:
    if not seconds or seconds <= 0:
        return "—"
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    if h > 0:
        return f"{h} t {m} min"
    if m > 0:
        return f"{m} min {s} sek"
    return f"{s} sek"


def build_docx(job: TranscriptionJob) -> bytes:
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    heading = document.add_heading(job.title or job.original_filename, level=0)
    for run in heading.runs:
        run.font.size = Pt(20)

    info = document.add_paragraph()
    info.add_run("Oprindeligt filnavn: ").bold = True
    info.add_run(job.original_filename or "—")
    info.add_run("\nOptagelsens længde: ").bold = True
    info.add_run(_format_duration(job.duration_seconds))
    info.add_run("\nDato: ").bold = True
    info.add_run(job.created_at.strftime("%d-%m-%Y %H:%M") if job.created_at else "—")
    info.add_run("\nSprog: ").bold = True
    info.add_run(job.language or "da")

    document.add_paragraph()

    text = (job.document.edited_text if job.document and job.document.edited_text else "").strip()
    for paragraph in text.split("\n\n") or [""]:
        p = document.add_paragraph(paragraph.strip())
        p.paragraph_format.space_after = Pt(6)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
