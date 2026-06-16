"""Word-eksport af et AI-genereret dokument (summary/referat/next_steps/followup)."""
from __future__ import annotations

import io
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt

from ...models import TranscriptionJob
from ..summary import DOCUMENT_LABELS

_BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")


def _add_runs_with_bold(paragraph, text: str) -> None:
    """Konverter **fed** til rigtige bold runs i en Word-afsnit."""
    last = 0
    for match in _BOLD_RE.finditer(text):
        if match.start() > last:
            paragraph.add_run(text[last : match.start()])
        run = paragraph.add_run(match.group(1))
        run.bold = True
        last = match.end()
    if last < len(text):
        paragraph.add_run(text[last:])


def _is_heading_line(line: str, next_line: str | None) -> bool:
    """Heuristik: en kort linje uden trailing tegnsætning, efterfulgt af punktopstilling, er en overskrift."""
    stripped = line.strip()
    if len(stripped) < 2 or len(stripped) > 80:
        return False
    if stripped.endswith((".", ",", ":", "?", "!", "—", "–")):
        return False
    if next_line is None:
        return False
    nxt = next_line.strip()
    return nxt.startswith(("-", "*", "•"))


def build_summary_docx(job: TranscriptionJob, doc_type: str, content: str) -> bytes:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    label = DOCUMENT_LABELS.get(doc_type, "Dokument")
    document.add_heading(label, level=0)

    meta = document.add_paragraph()
    meta.add_run("Optagelse: ").bold = True
    meta.add_run(job.title or job.original_filename or "—")
    meta.add_run("\nDato: ").bold = True
    meta.add_run(
        (job.created_at or datetime.utcnow()).strftime("%d-%m-%Y %H:%M")
    )
    if job.duration_seconds and job.duration_seconds > 0:
        h = int(job.duration_seconds // 3600)
        m = int((job.duration_seconds % 3600) // 60)
        meta.add_run("\nLængde: ").bold = True
        meta.add_run(f"{h} t {m} min" if h else f"{m} min")

    document.add_paragraph()

    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            document.add_paragraph()
            i += 1
            continue

        # Markdown-style overskrifter
        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        # Punktopstilling
        if re.match(r"^[-*•]\s+", stripped):
            text = re.sub(r"^[-*•]\s+", "", stripped)
            p = document.add_paragraph(style="List Bullet")
            _add_runs_with_bold(p, text)
            i += 1
            continue

        if re.match(r"^\d+[\.\)]\s+", stripped):
            text = re.sub(r"^\d+[\.\)]\s+", "", stripped)
            p = document.add_paragraph(style="List Number")
            _add_runs_with_bold(p, text)
            i += 1
            continue

        # Heuristik for overskrifter (kort linje før punktopstilling)
        next_line = lines[i + 1] if i + 1 < len(lines) else None
        if _is_heading_line(line, next_line):
            document.add_heading(stripped, level=2)
            i += 1
            continue

        # Almindeligt afsnit
        p = document.add_paragraph()
        _add_runs_with_bold(p, stripped)
        p.paragraph_format.space_after = Pt(6)
        i += 1

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
